"""Generate ReviewTrust Proposal Word Document.

Đồng bộ với docs/00_proposal.md (phiên bản sau khi rà soát theo góp ý GVHD).
"""
from __future__ import annotations

import os
from typing import Sequence

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

doc = Document()

# ============================================================
# Page setup
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(13)
normal.paragraph_format.line_spacing = 1.5


# ============================================================
# Helpers
# ============================================================
def h(text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)


def p(
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    size: int = 13,
) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        para.alignment = align
    para.paragraph_format.space_after = Pt(4)


def bullets(items: Sequence[str]) -> None:
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)


def table(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
    for row_data in rows:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_paragraph()


def page_break() -> None:
    doc.add_page_break()


# ============================================================
# TRANG BÌA
# ============================================================
for _ in range(3):
    doc.add_paragraph()
p("BỘ GIÁO DỤC VÀ ĐÀO TẠO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
p("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP.HCM (HUTECH)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
p("KHOA CÔNG NGHỆ THÔNG TIN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)

for _ in range(3):
    doc.add_paragraph()

p("ĐỀ CƯƠNG ĐỒ ÁN CHUYÊN NGÀNH", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run(
    "REVIEWTRUST (VoidRV)\n"
    "HỆ THỐNG XÁC ĐỊNH ĐỘ TIN CẬY REVIEW NHÀ HÀNG TIẾNG VIỆT\n"
    "DỰA TRÊN PHOBERT VÀ PHÂN TÍCH HÀNH VI REVIEWER"
)
run.font.name = "Times New Roman"
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

for _ in range(4):
    doc.add_paragraph()

for label, value in [
    ("Chuyên ngành:", "Hệ thống Thông tin Ứng dụng"),
    ("Sinh viên thực hiện:", "Trương Viết Hiệp"),
    ("MSSV:", "2380600637"),
    ("Giảng viên hướng dẫn:", "..............................................."),
]:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = para.add_run(f"{label} ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(13)
    r1.bold = True
    r2 = para.add_run(value)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

p("TP. Hồ Chí Minh, 2026", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()

# ============================================================
# MỤC LỤC
# ============================================================
h("MỤC LỤC", level=1)
for line in [
    "1. Tên đề tài",
    "2. Giới thiệu và đặt vấn đề",
    "   2.1. Thực trạng",
    "   2.2. Bài toán cần giải quyết",
    "   2.3. Phạm vi ứng dụng",
    "   2.4. Lý do chọn đề tài",
    "3. Hiện trạng bài toán và khảo sát nghiên cứu liên quan",
    "   3.1. Dòng chảy nghiên cứu Fake Review Detection",
    "   3.2. Các nghiên cứu quốc tế tiêu biểu (5 paper 2025–2026)",
    "   3.3. Các nghiên cứu tại Việt Nam",
    "   3.4. Bảng tổng hợp và phân tích hạn chế",
    "   3.5. Khoảng trống nghiên cứu (Research Gap)",
    "   3.6. Phương pháp kế thừa và hướng phát triển",
    "4. Nguồn dữ liệu và chiến lược xây dựng tập dữ liệu",
    "   4.1. Nguồn dữ liệu",
    "   4.2. Quy trình thu thập",
    "   4.3. Chiến lược gán nhãn",
    "   4.4. Đặc điểm tập dữ liệu",
    "5. Mục tiêu nghiên cứu",
    "   5.1. Mục tiêu chung",
    "   5.2. Mục tiêu cụ thể",
    "   5.3. Câu hỏi nghiên cứu",
    "6. Phân tích và thiết kế hệ thống (UC/ERD)",
    "   6.1. Phân rã chức năng",
    "   6.2. Sơ đồ Use Case",
    "   6.3. Đặc tả các Use Case chính",
    "   6.4. Sơ đồ ERD",
    "7. Phương pháp đề xuất — Kiến trúc 2 Layer",
    "8. Công nghệ sử dụng",
    "9. Kế hoạch thực hiện (12 tuần)",
    "10. Dự kiến rủi ro và giải pháp",
    "11. Cấu trúc quyển báo cáo dự kiến",
    "12. Tài liệu tham khảo",
]:
    p(line)
page_break()


# ============================================================
# 1. TÊN ĐỀ TÀI
# ============================================================
h("1. TÊN ĐỀ TÀI", level=1)
p(
    "ReviewTrust: Hệ thống xác định độ tin cậy review nhà hàng tiếng Việt "
    "kết hợp PhoBERT và phân tích hành vi reviewer",
    bold=True,
)

# ============================================================
# 2. GIỚI THIỆU VÀ ĐẶT VẤN ĐỀ
# ============================================================
h("2. GIỚI THIỆU VÀ ĐẶT VẤN ĐỀ", level=1)

h("2.1. Thực trạng", level=2)
p(
    "Ngành F&B (Food & Beverage) tại Việt Nam đang phát triển mạnh mẽ cùng với sự bùng nổ "
    "của du lịch tự túc. Các nền tảng đánh giá trực tuyến như Google Maps, Foody.vn, "
    "TripAdvisor đã trở thành kênh tham khảo chủ đạo của thực khách trước khi lựa chọn "
    "nhà hàng. Tuy nhiên, hệ sinh thái review hiện đang bị thao túng nghiêm trọng bởi "
    "nhiều vấn đề:"
)
bullets([
    "Review farm: dịch vụ “tăng sao”, “đẩy top” diễn ra công khai; chỉ vài trăm nghìn đồng, "
    "chủ quán có thể mua hàng trăm review 5 sao kèm bình luận có cánh.",
    "Review bombing: chiến dịch đánh 1 sao hàng loạt vì mâu thuẫn không liên quan chất "
    "lượng món ăn, phá hoại danh tiếng doanh nghiệp.",
    "Ghost accounts / tài khoản rác: tài khoản vô danh được tạo hàng loạt nhằm farm rating.",
    "Review do AI sinh (GAI-generated): từ 2023 trở đi, số lượng fake review sinh bởi "
    "ChatGPT/LLM tăng đột biến; Luo et al. (2026 — Decision Support Systems) chỉ ra đây "
    "là mối đe doạ mới nổi cấp bách.",
    "Review generic: “Ngon lắm!”, “Tuyệt vời!”… — không có giá trị tham khảo nhưng vẫn "
    "tính vào rating trung bình.",
])
p(
    "Theo báo cáo của chính phủ Anh (2023) được Xu & Huo (2026) trích dẫn, fake review gây "
    "thiệt hại cho người tiêu dùng £500 triệu – £3,1 tỷ/năm. TripAdvisor đã phát hiện "
    "1,3 triệu fake review chỉ trong năm 2022."
)

h("2.2. Bài toán cần giải quyết", level=2)
p(
    "Cho một review nhà hàng gồm (văn bản tiếng Việt + số sao + metadata reviewer + "
    "thời điểm đăng), hệ thống phải đưa ra: (1) Trust Score định lượng trong thang 0–100; "
    "(2) Nhãn phân loại “Đáng tin cậy” / “Cần thận trọng” / “Nghi ngờ”; (3) Danh sách lý "
    "do giải thích (explainable) cho quyết định. Đồng thời, ở mức nhà hàng, hệ thống "
    "tổng hợp các Trust Score review thành Adjusted Rating (rating thật sau khi lọc "
    "review đáng ngờ).",
    italic=True,
)
p(
    "Đây là bài toán phân loại nhị phân (fake/genuine) có gắn điểm tin cậy liên tục và có "
    "lời giải thích — nằm ở giao điểm của ba lĩnh vực: Fake Review Detection (FRD), NLP "
    "tiếng Việt, và Explainable AI."
)

h("2.3. Phạm vi ứng dụng", level=2)
p("Phạm vi thực hiện:", bold=True)
bullets([
    "Web application (React 18 + FastAPI) với 2 trang: phân tích review đơn lẻ (demo) và dashboard tổng hợp quán.",
    "Fine-tune PhoBERT-base (VinAI) cho binary classification (genuine vs fake) trên dữ liệu F&B tiếng Việt.",
    "Rule-based behavior scoring dựa trên metadata reviewer (review count, frequency, burst, rating pattern).",
    "Scrape dữ liệu từ Google Maps (chính) và Foody.vn (bổ sung) để xây dựng dataset và demo.",
    "Copy-paste detection bằng SimHash.",
    "Ngôn ngữ: tiếng Việt. Miền: nhà hàng / quán ăn / quán nước tại Việt Nam.",
])
p("Ngoài phạm vi (KHÔNG thực hiện):", bold=True)
bullets([
    "Không xử lý ảnh, GPS, Vision Transformer (ViT).",
    "Không eKYC, không nhận dạng danh tính người dùng.",
    "Không phát triển browser extension hay mobile application.",
    "Không train model từ đầu — chỉ fine-tune pretrained PhoBERT.",
    "Không kiểm duyệt tự động hay báo cáo vi phạm cho nền tảng.",
    "Không mở rộng sang khách sạn, e-commerce, TripAdvisor, Yelp.",
])
table(
    ["Trục phạm vi", "Trong phạm vi", "Ngoài phạm vi"],
    [
        ["Ngôn ngữ", "Tiếng Việt", "Tiếng Anh, đa ngôn ngữ"],
        ["Miền", "Nhà hàng / quán ăn / quán nước", "Khách sạn, E-commerce"],
        ["Nguồn dữ liệu", "Google Maps (chính); Foody.vn (bổ sung)", "TripAdvisor, Yelp, Facebook"],
        ["Modality", "Văn bản + metadata số", "Ảnh, Video, GPS"],
        ["Hình thức sản phẩm", "Web Application (SPA)", "Mobile app, browser extension"],
        ["ML approach", "Fine-tune pretrained PhoBERT", "Train từ đầu, LLM API runtime"],
    ],
)

h("2.4. Lý do chọn đề tài", level=2)
bullets([
    "Tính cấp thiết thực tiễn: Người tiêu dùng Việt Nam hiện không có công cụ nào giúp "
    "đánh giá độ tin cậy của từng review Google Maps — một “nỗi đau” cụ thể, rõ ràng.",
    "Khoảng trống nghiên cứu: Các nghiên cứu FRD tiếng Việt hiện có (nhóm UIT 2022, "
    "2024) đều tập trung miền e-commerce (Tiki, Shopee). Miền F&B tiếng Việt với ngữ "
    "cảnh địa lý, cảm tính cao vẫn là một khoảng trống (xem mục 3).",
    "Khả thi về kỹ thuật: PhoBERT (VinAI) đã được pre-train mạnh cho tiếng Việt, "
    "fine-tune cho binary classification là khả thi với cấu hình cá nhân (RTX 3060 12GB).",
    "Phù hợp chuyên ngành Hệ thống Thông tin Ứng dụng: đồ án tích hợp ML pipeline + Web "
    "application + Database + API design — đúng chuyên ngành HTTTUD.",
    "Khả thi về dữ liệu: Google Maps là nguồn công khai, có thể scrape hợp pháp bằng "
    "Playwright ở quy mô nhỏ cho mục đích nghiên cứu học thuật.",
])
page_break()


# ============================================================
# 3. HIỆN TRẠNG BÀI TOÁN VÀ KHẢO SÁT NGHIÊN CỨU LIÊN QUAN
# ============================================================
h("3. HIỆN TRẠNG BÀI TOÁN VÀ KHẢO SÁT NGHIÊN CỨU LIÊN QUAN", level=1)

h("3.1. Phân loại nhóm phương pháp và dòng chảy nghiên cứu", level=2)
p(
    "Để xác định khoảng trống một cách có hệ thống, nghiên cứu được khảo sát theo "
    "5 nhóm phương pháp chính:"
)
p("Nhóm 1 — Linguistic + ML cổ điển (SVM, LR, Naive Bayes):", bold=True)
p(
    "Khai thác đặc trưng ngôn ngữ học thủ công: TF-IDF, n-gram, POS tag, độ dài câu. "
    "Đại diện: Ott et al. (2011) đạt 89,8% accuracy trên dataset hotel deception."
)
table(
    ["Ưu điểm", "Hạn chế với bài toán"],
    [
        ["Đơn giản, giải thích được; inference nhanh", "Không capture ngữ nghĩa sâu; F1 giảm khi domain shift; không có behavioral; không phù hợp tiếng Việt"],
    ],
)
p("Nhóm 2 — Graph-based / Network-based:", bold=True)
p(
    "Mô hình mạng lưới reviewer–review–sản phẩm. Đại diện: SpEagle (Rayana & Akoglu 2015) "
    "dùng belief propagation trên heterogeneous graph."
)
table(
    ["Ưu điểm", "Hạn chế"],
    [
        ["Phát hiện cụm review farm; semi-supervised", "Cần network lớn; không khai thác text; khó tích hợp web app"],
    ],
)
p("Nhóm 3 — Deep Learning tuần tự (CNN, LSTM, BiLSTM):", bold=True)
p(
    "Học đặc trưng tự động từ chuỗi text. Đại diện: NetSpam (Shehnepoor et al., 2017)."
)
table(
    ["Ưu điểm", "Hạn chế"],
    [
        ["Không cần feature engineering; capture sequential pattern", "Không có pre-training mạnh cho tiếng Việt; chỉ text; hiệu quả thấp hơn transformer"],
    ],
)
p("Nhóm 4 — Pre-trained Language Models (BERT, RoBERTa, PhoBERT):", bold=True)
p(
    "Fine-tune transformer. Đại diện: DeceptiveBERT (Bhatt 2022); UIT Vietnamese "
    "(Dinh & Luu 2022, 2024) dùng PhoBERT cho e-commerce."
)
table(
    ["Ưu điểm", "Hạn chế"],
    [
        ["SOTA semantic; PhoBERT phù hợp tiếng Việt; transfer learning mạnh", "Content-only; VN chỉ có e-commerce, chưa F&B; thiếu XAI"],
    ],
)
p("Nhóm 5 — Hybrid / Multimodal / LLM-aware (trọng tâm, 2022–2026):", bold=True)
p(
    "Kết hợp text + behavioral metadata + ảnh hoặc LLM phát hiện AI-generated review. "
    "Đại diện: ConvRoBERTa (Mewada 2026), IFML (Xu & Huo 2026), Luo et al. (2026)."
)
table(
    ["Ưu điểm", "Hạn chế"],
    [
        ["Accuracy cao nhất; tích hợp đa nguồn; có XAI", "Chủ yếu tiếng Anh; cần ảnh hoặc LLM API; chưa có F&B VN"],
    ],
)
p("Tổng hợp theo thời gian:", bold=True)
table(
    ["Giai đoạn", "Nhóm phương pháp", "F1 điển hình", "Khoảng trống"],
    [
        ["2008–2013", "Nhóm 1 (Linguistic + ML)", "0,65–0,80", "Không behavioral; không deep learning"],
        ["2014–2017", "Nhóm 2 (Graph) + Nhóm 3 (DL)", "0,78–0,89", "Chưa dùng pre-trained LM"],
        ["2018–2022", "Nhóm 4 (BERT/RoBERTa)", "0,84–0,90", "Content-only; không F&B VN; không XAI"],
        ["2023–2026", "Nhóm 5 (Hybrid/LLM/Multimodal)", "0,88–0,94", "Tiếng Anh; cần ảnh/LLM API; chưa có cho tiếng Việt F&B"],
    ],
)

h("3.2. Các nghiên cứu quốc tế tiêu biểu (khảo sát 5 paper 2025–2026)", level=2)

h("(1) Luo, Nan & Li (2026) — AI-generated Fake Review Detection", level=3)
p("Decision Support Systems (2026).", italic=True)
bullets([
    "Vấn đề: LLM (ChatGPT, Gemini) sinh fake review hàng loạt, khó phân biệt bằng mắt thường.",
    "Phương pháp: Xây dựng 2 loại biến (linguistic: perplexity, burstiness; emotional: "
    "cường độ cảm xúc) → áp dụng Outlier Detection dựa trên Cumulative Probability "
    "Density → huấn luyện AdaBoost phân loại.",
    "Kết quả: Vượt nhiều baseline state-of-the-art trên dữ liệu e-commerce.",
    "Hạn chế: Chỉ thực nghiệm tiếng Anh; cần dữ liệu có nhãn; chưa xét yếu tố behavior.",
])

h("(2) Zhang, Ngai, Xia & Wu (2025) — DCOC: Dynamic Classification for Online Content", level=3)
p("Knowledge-Based Systems (2025).", italic=True)
bullets([
    "Vấn đề: Dữ liệu review là dòng stream liên tục, thay đổi theo thời gian, nhiễu cao.",
    "Phương pháp: Kernel-based online learning với slope-adjusted ramp loss chống nhiễu; "
    "học incremental không cần retrain toàn bộ.",
    "Kết quả: Giữ độ chính xác ngay cả khi 30% dữ liệu bị nhiễu; thực nghiệm trên "
    "TripAdvisor, Yelp, Amazon.",
    "Hạn chế: Không phân tích nội dung sâu bằng transformer; phụ thuộc kernel chọn thủ công.",
])

h("(3) Qi, Li, Yang & Li (2025) — Transfer Learning for Fake News Detection (Survey)", level=3)
p("Information Fusion (2025).", italic=True)
bullets([
    "Đóng góp: Hệ thống hoá transfer learning trong FRD thành 3 loại — cross-domain, "
    "domain adaptation, domain generalization.",
    "Phát hiện quan trọng: (i) Lớp nhãn mất cân bằng (fake:genuine ≈ 1:300); (ii) Seesaw "
    "effect khi transfer cross-domain; (iii) Low-resource languages như tiếng Việt gặp "
    "degradation mạnh khi dùng model train trên tiếng Anh/Trung.",
    "Gợi ý: Với bài toán FRD cho ngôn ngữ ít tài nguyên → phải fine-tune lại trên dữ "
    "liệu bản địa, không thể dùng model ngoại ngữ trực tiếp.",
])

h("(4) Xu & Huo (2026) — IFML: Interpretable FRD đa phương thức + Human-Computer", level=3)
p("Applied Soft Computing (2026).", italic=True)
bullets([
    "Phương pháp: (a) Multimodal VAE (MVAE) fuse text + ảnh + đặc trưng cấu trúc; (b) "
    "LLM sinh giải thích tự nhiên; (c) Bayesian framework kết hợp quyết định máy + "
    "feedback người dùng.",
    "Kết quả: PR AUC 94,43% trên 3 bộ Yelp; tăng 2,55% so với baseline tốt nhất.",
    "Điểm đáng học: (i) Interpretability là bắt buộc — người dùng cần “vì sao”, không chỉ "
    "“fake/genuine”; (ii) contextual details càng phong phú → xác suất fake càng thấp; "
    "(iii) human-in-the-loop tăng độ chính xác đáng kể.",
    "Hạn chế: Yêu cầu dữ liệu đa phương thức (ảnh) không phải lúc nào cũng có; LLM API "
    "tốn chi phí runtime.",
])

h("(5) Mewada & Dewang (2026) — ConvRoBERTa: Fusing Sequential + Weighted Nonsequential", level=3)
p("IEEE Transactions on Artificial Intelligence (2026).", italic=True)
bullets([
    "Vấn đề cốt lõi: Các model trước (CNN-LSTM, DeceptiveBERT) chỉ dùng linguistic, bỏ "
    "qua behavioral; hoặc gán trọng số bằng nhau cho mọi đặc trưng.",
    "Phương pháp: (a) CART tính feature importance cho đặc trưng nonsequential "
    "(behavior metadata); (b) CNN trên các feature đã weighted; (c) RoBERTa encode text; "
    "(d) Scaled dot-product attention fuse cả hai.",
    "Kết quả: Accuracy 91,93% trên Yelp (ConvRoBERTa-SVM), vượt baseline 2,94%.",
    "Điểm đáng học (critical): Đây là phương pháp mạnh nhất và gần nhất với bài toán "
    "của chúng ta — fuse sequential + nonsequential giống 2-layer architecture; có cơ "
    "chế weighting cho feature quan trọng; phát hiện burstiness.",
    "Hạn chế áp dụng trực tiếp: Dùng RoBERTa (tiếng Anh); bỏ qua giải thích tự nhiên; "
    "chưa kiểm tra trên tiếng Việt / F&B.",
])

h("3.3. Các nghiên cứu tại Việt Nam", level=2)
table(
    ["Năm", "Tác giả", "Miền", "Phương pháp", "Kết quả", "Hạn chế"],
    [
        ["2022", "Dinh, Luu et al. (UIT)", "E-commerce (Shopee, Tiki, Lazada)",
         "PhoBERT fine-tune + rules", "F1 ≈ 0,78",
         "Chỉ text, không behavior; không phải F&B"],
        ["2024", "Dinh, Luu (UIT)", "E-commerce",
         "PhoBERT + metadata integration", "F1 ≈ 0,82",
         "Vẫn chỉ e-commerce; metadata đơn giản"],
        ["2020", "Nguyen & Nguyen (VinAI)", "Nền tảng NLP",
         "Pre-train PhoBERT-base 20GB", "SOTA NER, POS, DP",
         "Không dành riêng cho FRD"],
    ],
)

h("3.4. Bảng tổng hợp và phân tích hạn chế", level=2)
table(
    ["Phương pháp", "Năm", "Đặc trưng", "F1 / Acc", "Hạn chế với bài toán"],
    [
        ["Jindal & Liu (LR + duplicate)", "2008", "Linguistic only", "~0,65", "Không bắt được spam tinh vi"],
        ["Ott et al. (SVM + n-gram)", "2011", "Linguistic", "~0,86", "Chỉ domain hotel có crowd-sourced"],
        ["NetSpam (heterogeneous graph)", "2017", "Graph + metadata", "~0,89", "Yêu cầu review network lớn"],
        ["DeceptiveBERT", "2022", "Contextual text", "~0,88", "Bỏ qua behavioral"],
        ["ConvRoBERTa (Mewada 2026)", "2026", "Text + behavior fuse + attention", "91,93%", "Tiếng Anh, chưa F&B VN"],
        ["IFML (Xu 2026)", "2026", "Multimodal + LLM + Bayesian", "94,43% PR-AUC", "Yêu cầu ảnh; runtime LLM đắt"],
        ["DCOC (Zhang 2025)", "2025", "Online streaming + robust loss", "~0,87", "Không dùng transformer"],
        ["Luo (AdaBoost + outlier)", "2026", "AI-generated detection", "SOTA", "Tiếng Anh, không behavior"],
        ["UIT Vietnamese (Dinh 2024)", "2024", "PhoBERT + metadata", "~0,82", "E-commerce, không F&B"],
    ],
)

h("3.5. Khoảng trống nghiên cứu và suy ra Mục tiêu", level=2)
p("Từ khảo sát 5 nhóm phương pháp, xác định 4 khoảng trống rõ ràng:")
bullets([
    "Gap 1 — Ngôn ngữ x Miền chưa được phủ: Chưa có công trình công bố mô hình FRD riêng "
    "cho tiếng Việt + miền F&B. Nhóm 4 (PhoBERT) chỉ có e-commerce; Nhóm 5 toàn tiếng Anh.",
    "Gap 2 — Content + Behavior còn hạn chế trong VN: Dinh (2024) có metadata nhưng đơn "
    "giản; Mewada (2026) có fuse tốt nhưng tiếng Anh. VN vẫn chủ yếu dừng ở content-only.",
    "Gap 3 — Thiếu XAI: Hầu hết model FRD tiếng Việt chỉ trả về label, thiếu lý do giải "
    "thích. Xu & Huo (2026) chứng minh XAI tăng trust người dùng nhưng chưa có cho VN.",
    "Gap 4 — Thiếu sản phẩm web cho end-user: Tất cả nhóm phương pháp dừng ở model, "
    "không có web app cho traveler dùng trực tiếp.",
])
p("Bảng ánh xạ Khoảng trống — Mục tiêu cụ thể:", bold=True)
table(
    ["Khoảng trống", "Mục tiêu giải quyết"],
    [
        ["Gap 1: Chưa có dataset F&B tiếng Việt có nhãn", "MT1 — Xây dựng data pipeline + dataset ~3.000 review"],
        ["Gap 1: Chưa có model FRD fine-tune cho F&B VN", "MT2 — Fine-tune PhoBERT-base trên dataset tự xây"],
        ["Gap 2: Content-only, thiếu behavioral features", "MT3, MT4 — Triển khai Content Module + Behavior Module"],
        ["Gap 2: Thiếu cơ chế fuse có trọng số tối ưu", "MT5 — Trust Engine 0,6 x Content + 0,4 x Behavior"],
        ["Gap 3: Không có XAI/giải thích", "MT5 — Mỗi review có explanation >= 3 lý do"],
        ["Gap 4: Không có sản phẩm web cho end-user", "MT6 — Web application React hoạt động end-to-end"],
        ["Validation đóng góp từng layer", "MT7 — Ablation study so với baseline"],
        ["Kiểm chứng thực tế", "MT8 — Case study 3-5 quán thật tại HCM"],
    ],
)

h("3.6. Phương pháp kế thừa và hướng phát triển", level=2)
p(
    "Phương pháp được kế thừa làm nền: ConvRoBERTa (Mewada & Dewang, 2026).",
    bold=True,
)
p("Lý do chọn:", bold=True)
bullets([
    "Là công trình 2026 mới nhất, trên tạp chí IEEE TAI uy tín.",
    "Kiến trúc fuse sequential (text) + nonsequential (behavior) khớp hoàn toàn với "
    "2-layer architecture chúng ta đang thiết kế.",
    "Có cơ chế feature weighting (CART) cho feature quan trọng — giải quyết đúng hạn chế "
    "“treat all features equally” mà chính paper chỉ ra.",
    "Kết quả 91,93% đủ mạnh để làm baseline đáng tin.",
])
p("Hướng phát triển (đóng góp mới của ReviewTrust):", bold=True)
table(
    ["Thành phần ConvRoBERTa", "Thay đổi / mở rộng của ReviewTrust"],
    [
        ["RoBERTa encoder", "PhoBERT-base (VinAI) cho tiếng Việt"],
        ["Dataset Yelp (tiếng Anh)", "Dataset tự xây: Google Maps + Foody miền F&B Việt Nam"],
        ["Weighting = CART", "Rule-based scoring + fine-tuned weight (đơn giản, dễ giải thích)"],
        ["Attention fusion", "Linear combination 0,6 × Content + 0,4 × Behavior (đủ cho scale đồ án)"],
        ["Không có giải thích", "XAI Explainer sinh danh sách lý do (kế thừa Xu & Huo 2026 nhưng dùng rule-based)"],
        ["Không có UI", "Web app React dashboard có Trust Gauge, Adjusted Rating, Suspicious Cluster"],
        ["Không xét copy-paste", "Bổ sung SimHash (kế thừa Rayana & Akoglu 2015)"],
        ["Không xét burst", "Bổ sung Burst detection (≥ 15 tài khoản mới cùng ngày → cờ đỏ)"],
    ],
)
p(
    "Tóm lại: Chúng ta kế thừa triết lý fuse Content + Behavior từ ConvRoBERTa, thay "
    "RoBERTa bằng PhoBERT cho tiếng Việt, thay attention phức tạp bằng linear "
    "combination + rules cho phù hợp quy mô đồ án chuyên ngành, và bổ sung XAI + Web "
    "app cho người dùng cuối.",
    italic=True,
)
page_break()


# ============================================================
# 4. NGUỒN DỮ LIỆU
# ============================================================
h("4. NGUỒN DỮ LIỆU VÀ CHIẾN LƯỢC XÂY DỰNG TẬP DỮ LIỆU", level=1)

h("4.1. Nguồn dữ liệu", level=2)
table(
    ["Nguồn", "Phương pháp thu thập", "Số lượng dự kiến", "Vai trò"],
    [
        ["Google Maps", "Playwright headless browser", "~1.500–2.000 review", "Tập chính — train/val/test"],
        ["Foody.vn", "BeautifulSoup + HTML parsing", "~500–800 review", "Tập bổ sung — tăng đa dạng văn phong"],
        ["LLM-generated (GPT/Gemini)", "Prompt tạo fake review theo template", "~400–600 review", "Augment fake class để cân bằng"],
        ["Dataset công khai (Ott 2011 / Yelp Chi 2013)", "Tải + dịch + rà soát", "~300 review", "Transfer cross-lingual ban đầu"],
    ],
)
p("Tổng dự kiến: 2.700–3.700 review có nhãn.", bold=True)

h("4.2. Quy trình thu thập", level=2)
p("Bước 1 — Chọn mẫu quán ăn:", bold=True)
bullets([
    "Chọn 30–50 nhà hàng tại TP.HCM và Hà Nội, đa dạng loại hình: phở, trà sữa, bún bò, "
    "buffet, cà phê, quán chay, quán nướng.",
    "Cấp độ: từ vỉa hè đến nhà hàng cao cấp (để có đa dạng mức rating).",
    "Chỉ chọn quán có ≥ 30 review để đảm bảo đủ dữ liệu phân tích hành vi.",
])
p("Bước 2 — Scrape Google Maps:", bold=True)
bullets([
    "Dùng Playwright + Chromium headless, random User-Agent, delay 2–5s giữa các request "
    "để tránh rate limit.",
    "Mỗi review thu thập: content, star_rating, reviewer_name, reviewer_review_count, "
    "reviewer_photo (optional), posted_at, language.",
    "Bỏ qua review không phải tiếng Việt.",
])
p("Bước 3 — Scrape Foody.vn:", bold=True)
bullets([
    "Dùng BeautifulSoup trên HTML tĩnh, không cần headless.",
    "Thu thập cùng schema để có thể join 2 nguồn.",
])
p("Bước 4 — Tiền xử lý:", bold=True)
bullets([
    "Làm sạch text: loại emoji thừa, URL, @mention, chuẩn hoá dấu câu, Unicode NFC.",
    "Word-segmentation bằng pyvi / underthesea (bắt buộc cho PhoBERT).",
    "Tính SimHash 64-bit cho mỗi review để phát hiện duplicate.",
])

h("4.3. Chiến lược gán nhãn (Data Labeling)", level=2)
p(
    "Đây là khâu quan trọng nhất vì dataset không có nhãn sẵn. Áp dụng chiến lược Hybrid "
    "Labeling 3 tầng:"
)
p("Tầng 1 — Heuristic Rules (gán nhãn tự động sơ bộ):", bold=True)
table(
    ["Luật", "Nhãn gợi ý"],
    [
        ["Review < 5 từ hoặc toàn emoji", "fake / low-quality"],
        ["Reviewer < 3 review đời + đăng trong 1h với review 5★ toàn quán cùng khu",
         "fake (farm pattern)"],
        ["SimHash similarity > 90% với ≥ 2 review khác trên platform", "fake (copy-paste)"],
        ["Sentiment tích cực mạnh đối lập rating ≤ 2★ (hoặc ngược lại)",
         "fake (suspicious)"],
        ["≥ 3 aspects + ≥ 50 từ + reviewer ≥ 10 review đa dạng", "genuine (strong)"],
        ["Nằm trong burst day (≥ 15 tài khoản mới cùng ngày, ≥ 5★)", "fake (burst)"],
    ],
)
p("Tầng 2 — LLM-assisted Labeling (Luo et al. 2026 gợi ý):", bold=True)
bullets([
    "Random sample ~30% dataset, dùng GPT-4 / Gemini với prompt chuyên biệt để cross-check.",
    "Nếu LLM khớp heuristic → giữ nhãn; nếu lệch → chuyển tầng 3 arbitration.",
])
p("Tầng 3 — Manual Review (người gán nhãn chuẩn):", bold=True)
p("Người thực hiện gán nhãn:", bold=True)
bullets([
    "Sinh viên thực hiện (Trương Viết Hiệp): gán nhãn độc lập ~500 review tập test và "
    "~300–500 review có mâu thuẫn từ tầng 1 & 2.",
    "Một sinh viên khác cùng khoa (annotator thứ 2): gán nhãn độc lập cùng 500 review "
    "tập test để tính inter-annotator agreement — không nhìn nhãn của annotator 1.",
    "Giảng viên hướng dẫn: kiểm tra ngẫu nhiên ~50–100 review borderline, xác nhận "
    "guideline, quyết định cuối khi 2 annotator bất đồng (arbitration).",
])
p("Quy trình:", bold=True)
bullets([
    "Mỗi review được gán nhãn độc lập bởi ≥ 2 người → gold label.",
    "Nếu 2 annotator bất đồng → đưa lên giảng viên hướng dẫn quyết định.",
    "Tính Cohen’s Kappa giữa 2 annotator trên 500 review tập test. Kỳ vọng κ ≥ 0,70 "
    "(substantial agreement). Kết quả κ sẽ báo cáo trong quyển đồ án.",
])

h("4.4. Đặc điểm và thống kê dự kiến", level=2)
table(
    ["Thuộc tính", "Giá trị dự kiến"],
    [
        ["Tổng số review", "~3.000"],
        ["Số quán", "30–50"],
        ["Phân bố nhãn", "fake ~35% / genuine ~65% (imbalanced nhẹ)"],
        ["Độ dài trung bình", "25–40 từ"],
        ["% review có aspect cụ thể", "~60%"],
        ["% review có reviewer metadata đầy đủ", "~90%"],
        ["Chia train/val/test", "70% / 15% / 15%"],
    ],
)
p(
    "Tập dữ liệu sẽ được công bố công khai tùy chính sách Google Maps TOS — nếu không "
    "thể public full text thì công bố dạng anonymized + hash.",
    italic=True,
)
page_break()


# ============================================================
# 5. MỤC TIÊU NGHIÊN CỨU
# ============================================================
h("5. MỤC TIÊU NGHIÊN CỨU", level=1)

h("5.1. Mục tiêu chung", level=2)
p(
    "Xây dựng và thực nghiệm một hệ thống web (Web Application) có khả năng tự động đánh "
    "giá độ tin cậy của review nhà hàng tiếng Việt bằng cách kết hợp phân tích nội dung "
    "(PhoBERT + rules) và phân tích hành vi reviewer, đồng thời cung cấp lời giải thích "
    "rõ ràng cho từng đánh giá và dashboard tổng hợp cho từng nhà hàng.",
    italic=True,
)

h("5.2. Mục tiêu cụ thể (phân rã từ mục tiêu chung)", level=2)
table(
    ["Mã", "Mục tiêu cụ thể", "KPI đo lường"],
    [
        ["MT1", "Xây dựng data pipeline thu thập + tiền xử lý + gán nhãn dataset F&B tiếng Việt",
         "~3.000 review có nhãn, κ ≥ 0,70, documentation rõ ràng"],
        ["MT2", "Fine-tune PhoBERT-base cho binary classification genuine vs fake",
         "F1-macro ≥ 0,83, Accuracy ≥ 0,85 trên test set"],
        ["MT3", "Triển khai Layer 1 — Content Module (PhoBERT + sentiment-star + aspect + TTR + SimHash)",
         "Unit test coverage ≥ 80%, kiểm thử 100+ case"],
        ["MT4", "Triển khai Layer 2 — Behavior Module (count + frequency + burst + rating pattern)",
         "Unit test coverage ≥ 80%"],
        ["MT5", "Thiết kế Trust Engine gộp 2 layer → Trust Score + badge + explanation",
         "Mỗi review có ≥ 3 lý do explanation, Trust Score 0–100"],
        ["MT6", "Xây dựng Web App 3 trang: /, /restaurant/:slug, /analyze",
         "SPA React hoạt động, UX test OK với ≥ 5 người dùng"],
        ["MT7", "Thực nghiệm Ablation Study chứng minh đóng góp từng layer",
         "Bảng: PhoBERT-only vs +Behavior vs full; tăng ≥ 3% F1"],
        ["MT8", "Case Study thực tế trên 3–5 quán tại HCM",
         "Phát hiện ≥ 1 cluster suspicious với evidence rõ ràng"],
    ],
)

h("5.3. Câu hỏi nghiên cứu (Research Questions)", level=2)
bullets([
    "RQ1: Việc kết hợp PhoBERT (content) với behavioral rules có cải thiện F1 bao nhiêu "
    "so với chỉ dùng PhoBERT trên dữ liệu F&B tiếng Việt?",
    "RQ2: Trong behavioral features, feature nào (count / frequency / burst / rating "
    "pattern) đóng góp nhiều nhất vào performance?",
    "RQ3: Cách trình bày XAI (danh sách lý do) có giúp người dùng cuối tin vào Trust "
    "Score hơn so với chỉ trả về nhãn không? (Khảo sát UX N=20).",
])
page_break()


# ============================================================
# 6. PHÂN TÍCH & THIẾT KẾ HỆ THỐNG (UC / ERD)
# ============================================================
h("6. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG (UC/ERD)", level=1)

h("6.1. Phân rã chức năng", level=2)
p("Phân rã mục tiêu thành các chức năng hệ thống:")
function_tree = """ReviewTrust System
├── [F1] Thu thập dữ liệu
│   ├── F1.1 Scrape Google Maps từ URL (Playwright)
│   ├── F1.2 Scrape Foody (BeautifulSoup)
│   └── F1.3 Nhận review nhập tay (demo)
├── [F2] Phân tích nội dung (Content Module)
│   ├── F2.1 PhoBERT genuine probability inference
│   ├── F2.2 Sentiment–star consistency check
│   ├── F2.3 Aspect extraction (rule-based)
│   ├── F2.4 TTR + Length check
│   └── F2.5 SimHash duplicate detection
├── [F3] Phân tích hành vi (Behavior Module)
│   ├── F3.1 Reviewer history check (count)
│   ├── F3.2 Review frequency check
│   ├── F3.3 Burst detection
│   └── F3.4 Rating pattern check
├── [F4] Trust Engine
│   ├── F4.1 Gộp Content + Behavior → Trust Score
│   ├── F4.2 Quyết định badge
│   └── F4.3 Sinh explanation
├── [F5] Hiển thị Dashboard
│   ├── F5.1 Trust Gauge + Void Score
│   ├── F5.2 Adjusted Rating
│   ├── F5.3 Timeline + Burst highlight
│   ├── F5.4 Suspicious Clusters
│   └── F5.5 Review list + breakdown
└── [F6] Quản trị dữ liệu
    ├── F6.1 Lưu Review + Score vào DB
    └── F6.2 Quản lý scrape jobs (async)"""
tree_para = doc.add_paragraph()
run = tree_para.add_run(function_tree)
run.font.name = "Consolas"
run.font.size = Pt(11)

h("6.2. Sơ đồ Use Case", level=2)
p("Actor chính:")
bullets([
    "Traveler (Người dùng cuối): người đọc review, cần đánh giá độ tin cậy trước khi đi ăn.",
    "System (hệ thống tự động): bot scrape + background worker.",
])
p("Các Use Case:")
table(
    ["Mã UC", "Tên Use Case", "Actor chính"],
    [
        ["UC1", "Phân tích 1 review đơn lẻ (demo)", "Traveler"],
        ["UC2", "Phân tích nhà hàng từ URL Google Maps", "Traveler"],
        ["UC3", "Xem dashboard quán", "Traveler"],
        ["UC4", "Xem breakdown + XAI explanation của 1 review", "Traveler"],
        ["UC5", "Xem suspicious cluster của quán", "Traveler"],
        ["UC6", "Xem timeline + burst detection", "Traveler"],
        ["UC7", "Scrape reviews Google Maps (include từ UC2)", "System"],
        ["UC8", "Poll trạng thái scrape job", "Traveler"],
        ["UC9", "Background scoring toàn bộ review", "System"],
        ["UC10", "Lưu DB + cache", "System"],
    ],
)

h("6.3. Đặc tả các Use Case chính", level=2)

p("UC1 — Phân tích 1 review (Demo mode):", bold=True)
table(
    ["Thuộc tính", "Chi tiết"],
    [
        ["Actor", "Traveler"],
        ["Tiền điều kiện", "Web app đã load, API backend chạy"],
        ["Luồng chính",
         "1. User truy cập /analyze\n"
         "2. Nhập review text + số sao (1–5) vào form\n"
         "3. (Tuỳ chọn) nhập review_count của reviewer\n"
         "4. Nhấn “Phân tích”\n"
         "5. FE gọi POST /api/v1/analyze\n"
         "6. BE chạy Content Module → trả Trust Score + badge + explanation\n"
         "7. FE hiển thị gauge + danh sách lý do"],
        ["Luồng phụ", "6a. Không có reviewer_count → chỉ chạy Content → caveat “không có behavior data”"],
        ["Hậu điều kiện", "Kết quả hiển thị, không lưu DB"],
    ],
)

p("UC2 — Phân tích nhà hàng từ URL Google Maps:", bold=True)
table(
    ["Thuộc tính", "Chi tiết"],
    [
        ["Actor", "Traveler"],
        ["Tiền điều kiện", "URL hợp lệ Google Maps"],
        ["Luồng chính",
         "1. User paste URL vào trang chủ\n"
         "2. FE gọi POST /api/v1/scrape → nhận 202 + job_id\n"
         "3. FE poll GET /api/v1/scrape/status/{job_id} mỗi 3s\n"
         "4. Background worker scrape reviews (Playwright)\n"
         "5. Mỗi review: chạy Content + Behavior → lưu DB\n"
         "6. Worker hoàn tất → status = success\n"
         "7. FE redirect /restaurant/:slug"],
        ["Luồng phụ (lỗi)", "4a. Captcha / rate limit → job status = failed, hiện lỗi"],
        ["Hậu điều kiện", "Dữ liệu quán + reviews + trust_scores đã lưu DB"],
    ],
)

p("UC3 — Xem dashboard quán:", bold=True)
table(
    ["Thuộc tính", "Chi tiết"],
    [
        ["Actor", "Traveler"],
        ["Tiền điều kiện", "Quán đã được scrape & scored"],
        ["Luồng chính",
         "1. User truy cập /restaurant/:slug\n"
         "2. FE gọi GET /api/v1/restaurant/{slug}\n"
         "3. BE trả: info quán + reviews + scores + aggregated stats\n"
         "4. FE render: Trust Gauge tổng, Adjusted Rating, Timeline, Suspicious Cluster, list reviews"],
    ],
)

h("6.4. Sơ đồ ERD", level=2)
p("Hệ thống gồm 4 bảng chính:")
table(
    ["Bảng", "Column chính", "Ràng buộc"],
    [
        ["restaurants",
         "id (PK), name, slug (UQ), google_place_id (UQ), google_maps_url, address, last_scraped_at, created_at",
         "slug UNIQUE → URL thân thiện"],
        ["reviews",
         "id (PK), restaurant_id (FK), content, star_rating, reviewer_name, reviewer_review_count, posted_at, simhash, source, created_at",
         "simhash INDEX; source ∈ {google_maps, foody, manual}"],
        ["trust_scores",
         "id (PK), review_id (FK UQ), content_score, behavior_score, trust_score, void_score, badge, explanation (JSONB), model_version, content_only, created_at",
         "review_id UNIQUE (1–1 với reviews)"],
        ["scrape_jobs",
         "id (PK), restaurant_id (FK), status, error_message, started_at, finished_at",
         "status ∈ {pending, running, success, failed}"],
    ],
)
p(
    "Mối quan hệ: restaurants 1–* reviews; reviews 1–1 trust_scores; restaurants 1–* "
    "scrape_jobs.",
    italic=True,
)
page_break()


# ============================================================
# 7. PHƯƠNG PHÁP ĐỀ XUẤT
# ============================================================
h("7. PHƯƠNG PHÁP ĐỀ XUẤT — KIẾN TRÚC 2 LAYER", level=1)

h("7.1. Tổng quan", level=2)
p(
    "Kế thừa triết lý fuse sequential + nonsequential của ConvRoBERTa (Mewada & Dewang, "
    "2026), ReviewTrust dùng kiến trúc 2-layer đơn giản hoá phù hợp quy mô đồ án. "
    "Input là 1 review, output là Trust Score + badge + explanation."
)

h("7.2. Layer 1 — Content Module (trọng số 60%)", level=2)
table(
    ["Yếu tố", "Điều kiện", "Điểm"],
    [
        ["Base score", "PhoBERT genuine_prob × 100", "(0–100)"],
        ["Sentiment–star mâu thuẫn", "Text khen + star 1–2, hoặc ngược lại", "−25"],
        ["Aspect count", "≥ 3", "+15"],
        ["Aspect count", "= 2", "+10"],
        ["Aspect count", "= 0", "−15"],
        ["TTR", "< 0,4 (lặp từ nhiều)", "−10"],
        ["TTR", "> 0,7 + text dài", "+5"],
        ["Length", "< 10 từ", "−30"],
        ["Length", "10–19 từ", "−20"],
        ["Length", "≥ 50 từ + có aspect", "+5"],
        ["SimHash similarity", "> 90% với review khác", "−40"],
        ["SimHash similarity", "80–90%", "−30"],
    ],
)
p("Kết quả cuối cùng của Layer 1 được clamp về [0, 100].", italic=True)

h("7.3. Layer 2 — Behavior Module (trọng số 40%)", level=2)
table(
    ["Yếu tố", "Điều kiện", "Điểm"],
    [
        ["Review count", "< 3", "−15"],
        ["Review count", "3–4", "−10"],
        ["Review count", "≥ 10", "+5"],
        ["Frequency", "≥ 6 review trong 1h", "−30"],
        ["Burst", "Reviewer thuộc cụm ≥ 15 tài khoản mới cùng ngày", "−25"],
        ["Rating pattern", "Liên tiếp cùng sao (chỉ 5★ hoặc chỉ 1★)", "−15"],
    ],
)

h("7.4. Trust Engine", level=2)
trust_code = (
    "Trust Score = 0,60 × Content Score + 0,40 × Behavior Score\n"
    "Void Score  = 100 − Trust Score\n\n"
    "Badge:\n"
    "  Trust ≥ 75   → “Đáng tin cậy”    (trusted, #22c55e)\n"
    "  Trust 50–74  → “Cần thận trọng”  (caution, #eab308)\n"
    "  Trust < 50   → “Nghi ngờ”        (suspicious, #ef4444)"
)
code_para = doc.add_paragraph()
run = code_para.add_run(trust_code)
run.font.name = "Consolas"
run.font.size = Pt(11)
p(
    "Nếu demo mode (không có reviewer metadata): Trust = Content, gắn cờ content_only = true.",
    italic=True,
)

h("7.5. So sánh baseline (dự kiến khi báo cáo)", level=2)
table(
    ["Mã", "Baseline", "Thành phần", "F1 dự kiến"],
    [
        ["B1", "PhoBERT-only", "Chỉ PhoBERT", "0,80"],
        ["B2", "Rules-only", "Chỉ rule Content + Behavior", "0,72"],
        ["B3", "PhoBERT + Content rules", "Chỉ Layer 1", "0,82"],
        ["B4", "Full ReviewTrust", "Layer 1 + Layer 2 + Trust Engine", "0,85 (mục tiêu)"],
        ["B5", "ConvRoBERTa (tham chiếu)", "Full paper (tiếng Anh, Yelp)", "0,92"],
    ],
)
page_break()


# ============================================================
# 8. CÔNG NGHỆ
# ============================================================
h("8. CÔNG NGHỆ SỬ DỤNG", level=1)
table(
    ["Tầng", "Công nghệ", "Lý do chọn"],
    [
        ["ML / NLP", "PyTorch 2.x, HuggingFace Transformers, PhoBERT-base",
         "PhoBERT là SOTA tiếng Việt (Nguyen & Nguyen 2020)"],
        ["NLP utilities", "underthesea / pyvi (word-seg), datasketch (SimHash)",
         "Chuẩn cho PhoBERT input"],
        ["Backend", "FastAPI (Python 3.11+), SQLAlchemy 2.x async + asyncpg",
         "Async cao, type-safe, Swagger auto-gen"],
        ["Scraping", "Playwright 1.x (Chromium headless), BeautifulSoup4",
         "Playwright xử lý DOM động (Google Maps SPA)"],
        ["Database", "PostgreSQL 16", "JSONB cho explanation, full-text, ổn định"],
        ["Migration", "Alembic", "Version DB schema"],
        ["Frontend", "React 18 + Vite + TailwindCSS + Recharts",
         "Dev nhanh, component-based, chart đẹp"],
        ["Deploy", "Docker Compose → Railway / Render free tier",
         "Reproducible, chi phí 0đ"],
        ["Testing", "pytest (backend), Vitest (frontend)", "Chuẩn ngành"],
    ],
)
page_break()


# ============================================================
# 9. KẾ HOẠCH
# ============================================================
h("9. KẾ HOẠCH THỰC HIỆN (12 TUẦN)", level=1)
table(
    ["Tuần", "Công việc", "Deliverable"],
    [
        ["1", "Hoàn tất đề cương + khảo sát 5 paper + setup repo + env",
         "Proposal + Notes khảo sát + Repo Git"],
        ["2", "Scrape Google Maps ~2.000 review; Foody ~800 review; tiền xử lý",
         "Raw dataset CSV/JSONL"],
        ["3", "Gán nhãn tầng 1 (heuristics) + tầng 2 (LLM-assisted)",
         "Dataset labeled v0.5"],
        ["4", "Gán nhãn tầng 3 (manual 500 gold + 500 arbitration); tính Cohen’s Kappa",
         "Dataset final + κ report"],
        ["5", "Fine-tune PhoBERT (3–5 epochs) + hyper-param tuning",
         "phobert_voidrv.pt + metrics"],
        ["6", "Triển khai Content Module + Behavior Module + unit tests",
         "services/*.py + tests coverage 80%"],
        ["7", "Triển khai Trust Engine + XAI Explainer + integration tests",
         "API /analyze hoạt động"],
        ["8", "Scraper Google Maps + Background worker + scrape_jobs flow",
         "API /scrape + /scrape/status"],
        ["9", "Frontend 3 trang (Home + Analyze + Restaurant Dashboard)",
         "SPA React hoạt động end-to-end"],
        ["10", "Ablation study + Case study 3–5 quán thật + viết chương 4",
         "Notebook + bảng metrics + case study"],
        ["11", "Khảo sát UX N=20 + XAI user test + hoàn thiện UI", "UX report"],
        ["12", "Deploy Docker, viết báo cáo quyển, slide demo",
         "URL public + .docx + .pptx"],
    ],
)
page_break()


# ============================================================
# 10. RỦI RO
# ============================================================
h("10. DỰ KIẾN RỦI RO VÀ GIẢI PHÁP", level=1)
table(
    ["Rủi ro", "Mức độ", "Giải pháp phòng ngừa"],
    [
        ["Google Maps chặn scrape (CAPTCHA, rate-limit)", "Cao",
         "Random User-Agent, delay 2–5s, giới hạn ≤ 200 req/giờ, fallback dataset thủ công"],
        ["Dataset mất cân bằng nhãn nặng", "Cao",
         "Augment fake class bằng LLM (Luo 2026); class_weight + focal loss"],
        ["κ gán nhãn < 0,70 → nhãn không tin cậy", "Trung bình",
         "Bổ sung vòng arbitration thứ 2; rà guideline chặt hơn"],
        ["F1 PhoBERT < 0,80 sau fine-tune", "Trung bình",
         "Thử PhoBERT-large; back-translation augment; adapt ConvRoBERTa"],
        ["RTX 3060 12GB không đủ train PhoBERT-large", "Trung bình",
         "Gradient accumulation; batch_size=8; Colab Pro"],
        ["PhoBERT size ~540MB khó deploy free tier", "Thấp",
         "Export ONNX quantized INT8 (giảm ~75%); async inference queue"],
        ["Foody.vn đổi layout HTML", "Thấp",
         "Viết test e2e scraper; fallback chỉ dùng Google Maps"],
        ["Tranh chấp pháp lý khi scrape", "Thấp",
         "Chỉ dùng cho học thuật, không public raw data, ẩn danh reviewer_name trước khi công bố"],
    ],
)
page_break()


# ============================================================
# 11. CẤU TRÚC QUYỂN BÁO CÁO
# ============================================================
h("11. CẤU TRÚC QUYỂN BÁO CÁO DỰ KIẾN", level=1)
bullets([
    "MỞ ĐẦU: Tính cấp thiết, mục tiêu, phạm vi, phương pháp, ý nghĩa.",
    "CHƯƠNG 1 — TỔNG QUAN: Bài toán FRD, thực trạng, khảo sát nghiên cứu liên quan "
    "(5 paper), khoảng trống và hướng kế thừa.",
    "CHƯƠNG 2 — CƠ SỞ LÝ THUYẾT: PhoBERT, Transformer, SimHash, TTR, Aspect-based "
    "Sentiment, Behavioral Features, XAI.",
    "CHƯƠNG 3 — PHÂN TÍCH & THIẾT KẾ HỆ THỐNG: UC, ERD, API specs, kiến trúc 2-layer, "
    "luồng xử lý.",
    "CHƯƠNG 4 — CÀI ĐẶT & THỰC NGHIỆM: Fine-tune PhoBERT, ablation study, case study, "
    "UX test, so sánh baseline.",
    "CHƯƠNG 5 — GIAO DIỆN & DEMO: Screenshots web app, demo video, hướng dẫn sử dụng.",
    "KẾT LUẬN & HƯỚNG PHÁT TRIỂN: đóng góp, hạn chế, hướng mở rộng (Foody cross-"
    "platform, LLM-generated detection, stylometry).",
    "TÀI LIỆU THAM KHẢO.",
])
page_break()


# ============================================================
# 12. TÀI LIỆU THAM KHẢO
# ============================================================
h("12. TÀI LIỆU THAM KHẢO", level=1)

p("Quốc tế (theo tài liệu thầy gửi):", bold=True)
for ref in [
    "[1] Luo, J., Nan, G., & Li, D. (2026). AI-generated fake review detection. "
    "Decision Support Systems. Elsevier.",
    "[2] Zhang, Z., Ngai, E. W. T., Xia, S., & Wu, Z. (2025). Enhancing fake review "
    "detection: A robust and adaptive approach for data streams. Knowledge-Based Systems. Elsevier.",
    "[3] Qi, C., Li, X., Yang, X., & Li, Z. (2025). A review of fake news detection "
    "based on transfer learning. Information Fusion. Elsevier.",
    "[4] Xu, T., & Huo, Y. (2026). Interpretable fake review detection based on "
    "multimodal information and human–computer collaboration. Applied Soft Computing. Elsevier.",
    "[5] Mewada, A., & Dewang, R. K. (2026). ConvRoBERTa: Detecting Fake Reviews by "
    "Fusing Sequential and Weighted Nonsequential Features. IEEE Transactions on "
    "Artificial Intelligence, 7(3), 1301–1314. [Phương pháp nền kế thừa chính]",
]:
    p(ref)

p("Các nghiên cứu tiếng Việt:", bold=True)
for ref in [
    "[6] Dinh, C. V., Luu, S. T., et al. (2022). Detecting Spam Reviews on Vietnamese "
    "E-commerce Websites. arXiv:2207.14636.",
    "[7] Dinh, C. V., & Luu, S. T. (2024). Metadata Integration for Spam Reviews "
    "Detection on Vietnamese E-commerce Websites. arXiv:2405.13292.",
    "[8] Nguyen, D. Q., & Nguyen, A. T. (2020). PhoBERT: Pre-trained language models "
    "for Vietnamese. Findings of EMNLP 2020.",
]:
    p(ref)

p("Các công trình nền tảng:", bold=True)
for ref in [
    "[9] Jindal, N., & Liu, B. (2008). Opinion Spam and Analysis. WSDM.",
    "[10] Ott, M., Choi, Y., Cardie, C., & Hancock, J. T. (2011). Finding Deceptive "
    "Opinion Spam by Any Stretch of the Imagination. ACL.",
    "[11] Mukherjee, A., Venkataraman, V., Liu, B., & Glance, N. (2013). What Yelp Fake "
    "Review Filter Might Be Doing? ICWSM.",
    "[12] Rayana, S., & Akoglu, L. (2015). Collective Opinion Spam Detection. KDD.",
    "[13] Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training "
    "of Deep Bidirectional Transformers for Language Understanding. NAACL-HLT.",
]:
    p(ref)

p("Tài liệu kỹ thuật:", bold=True)
for ref in [
    "[14] FastAPI documentation — https://fastapi.tiangolo.com/",
    "[15] HuggingFace Transformers documentation — https://huggingface.co/docs/transformers",
    "[16] Playwright for Python documentation — https://playwright.dev/python/",
    "[17] PhoBERT GitHub (VinAIResearch) — https://github.com/VinAIResearch/PhoBERT",
]:
    p(ref)


# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), "ReviewTrust_Proposal.docx")
doc.save(output_path)
print(f"Saved to {output_path}")
